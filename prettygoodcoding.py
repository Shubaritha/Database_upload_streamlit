import streamlit as st
import psycopg2
import pandas as pd
import io
import base64
from docx import Document
import PyPDF2
import tempfile
import os
import tiktoken
import openai
import uuid
import numpy as np
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Get OpenAI API key from environment variables
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    st.error("OpenAI API key not found. Please make sure it's set in the .env file.")

st.set_page_config(page_title="Document to Neon Database Uploader", layout="wide")

st.title("Document to Neon Database Uploader")
st.write("Upload PDF or Word documents to your Neon Database")

# File upload section
st.header("Upload Document")
uploaded_file = st.file_uploader("Choose a PDF or Word document", type=['pdf', 'docx'])

# Database connection section
st.header("Database Connection")
connection_string = st.text_input("Enter your Neon Database connection string", 
                                 help="Format: postgresql://username:password@endpoint:port/dbname", 
                                 type="password")

# Define fixed max tokens value
MAX_TOKENS = 10000
st.info(f"Documents will be automatically split into chunks of maximum {MAX_TOKENS} tokens.")

# Function to extract text from PDF
def extract_text_from_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page_num in range(len(pdf_reader.pages)):
        page = pdf_reader.pages[page_num]
        text += page.extract_text()
    return text

# Function to extract text from DOCX
def extract_text_from_docx(file):
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp:
        temp.write(file.getvalue())
        temp_path = temp.name
    
    doc = Document(temp_path)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    
    # Clean up the temporary file
    os.unlink(temp_path)
    
    return text

# Function to check if a table exists and create it if it doesn't
def ensure_table_exists(conn_string, table_name):
    """Check if table exists and create it if it doesn't"""
    conn = None
    try:
        conn = psycopg2.connect(conn_string)
        cursor = conn.cursor()
        
        # Check if table exists - use lowercase for information_schema search
        cursor.execute(f"""
            SELECT EXISTS (
                SELECT FROM information_schema.tables 
                WHERE table_schema = 'public'
                AND table_name = '{table_name.lower()}'
            );
        """)
        table_exists = cursor.fetchone()[0]
        
        if not table_exists:
            # Create pgvector extension
            try:
                cursor.execute("CREATE EXTENSION IF NOT EXISTS vector;")
                conn.commit()
            except Exception as e:
                conn.rollback()
                st.warning(f"Could not create vector extension: {str(e)}")
            
            # Create the table - use double quotes to preserve case
            try:
                cursor.execute(f"""
                    CREATE TABLE "{table_name}" (
                        id SERIAL PRIMARY KEY,
                        chunk_id VARCHAR(255),
                        title VARCHAR(255),
                        contents TEXT,
                        vector vector(1536)
                    )
                """)
                conn.commit()
                st.success(f"Created new table '{table_name}'")
                return True
            except Exception as e:
                conn.rollback()
                st.error(f"Error creating table: {str(e)}")
                return False
        
        return True
    except Exception as e:
        st.error(f"Error checking/creating table: {str(e)}")
        return False
    finally:
        if conn is not None:
            conn.close()

# Function to ensure table has required columns for vector storage
def ensure_table_has_required_columns(conn_string, table_name):
    """Check if table has required columns and add them if they don't exist"""
    # First make sure the table exists
    if not ensure_table_exists(conn_string, table_name):
        return False
        
    conn = None
    try:
        conn = psycopg2.connect(conn_string)
        cursor = conn.cursor()
        
        # Create pgvector extension if it doesn't exist
        try:
            cursor.execute("CREATE EXTENSION IF NOT EXISTS vector;")
            conn.commit()
        except Exception as e:
            conn.rollback()
            st.warning(f"Note: Could not create vector extension. If it doesn't exist, vectors won't work: {str(e)}")
        
        # Get existing columns - table_name is used as-is with information_schema
        cursor.execute(f"""
            SELECT column_name FROM information_schema.columns
            WHERE table_name = '{table_name.lower()}'
            AND table_schema = 'public'
        """)
        existing_columns = [col[0] for col in cursor.fetchall()]
        
        # Define required columns and their types
        required_columns = {
            'id': 'SERIAL PRIMARY KEY',
            'chunk_id': 'VARCHAR(255)',
            'title': 'VARCHAR(255)',
            'contents': 'TEXT',
            'vector': 'vector(1536)'
        }
        
        # Check if each required column exists, if not add it
        for col_name, col_type in required_columns.items():
            if col_name not in existing_columns:
                try:
                    # Special handling for id column if it doesn't exist
                    if col_name == 'id':
                        try:
                            cursor.execute(f'ALTER TABLE "{table_name}" ADD COLUMN id SERIAL PRIMARY KEY')
                            conn.commit()
                        except Exception:
                            # If adding primary key fails, just add as serial
                            conn.rollback()  # Reset the aborted transaction
                            cursor.execute(f'ALTER TABLE "{table_name}" ADD COLUMN id SERIAL')
                            conn.commit()
                    else:
                        cursor.execute(f'ALTER TABLE "{table_name}" ADD COLUMN {col_name} {col_type}')
                        conn.commit()
                    st.info(f"Added column '{col_name}' to table '{table_name}'")
                except Exception as e:
                    conn.rollback()  # Reset the aborted transaction
                    st.warning(f"Could not add column '{col_name}': {str(e)}")
        
        cursor.close()
        return True
    except Exception as e:
        st.error(f"Error ensuring table has required columns: {str(e)}")
        return False
    finally:
        if conn is not None:
            conn.close()

# Function to count tokens
def count_tokens(text, model="cl100k_base"):
    """Count the number of tokens in the text using tiktoken"""
    try:
        encoder = tiktoken.get_encoding(model)
        return len(encoder.encode(text))
    except Exception as e:
        st.error(f"Error counting tokens: {str(e)}")
        return 0

# Function to split document into chunks
def split_document(text, filename, max_tokens=10000):
    """Split a document into chunks based on paragraphs and token count"""
    chunks = []
    paragraphs = text.split('\n\n')
    
    current_chunk = ""
    current_tokens = 0
    chunk_id = 1
    
    for paragraph in paragraphs:
        paragraph_tokens = count_tokens(paragraph)
        
        # If a single paragraph exceeds max_tokens, split it further
        if paragraph_tokens > max_tokens:
            words = paragraph.split()
            temp_para = ""
            for word in words:
                temp_para += word + " "
                if count_tokens(temp_para) >= max_tokens:
                    chunks.append({
                        "chunk_id": f"{filename}-{chunk_id}",
                        "title": f"{filename} (Part {chunk_id})",
                        "contents": temp_para.strip()
                    })
                    chunk_id += 1
                    temp_para = ""
            if temp_para:  # Add remaining text if any
                chunks.append({
                    "chunk_id": f"{filename}-{chunk_id}",
                    "title": f"{filename} (Part {chunk_id})",
                    "contents": temp_para.strip()
                })
                chunk_id += 1
        # Check if adding this paragraph would exceed the token limit
        elif current_tokens + paragraph_tokens > max_tokens and current_chunk:
            chunks.append({
                "chunk_id": f"{filename}-{chunk_id}",
                "title": f"{filename} (Part {chunk_id})",
                "contents": current_chunk.strip()
            })
            chunk_id += 1
            current_chunk = paragraph
            current_tokens = paragraph_tokens
        else:
            current_chunk += "\n\n" + paragraph if current_chunk else paragraph
            current_tokens += paragraph_tokens
    
    # Add the last chunk if it's not empty
    if current_chunk:
        chunks.append({
            "chunk_id": f"{filename}-{chunk_id}",
            "title": f"{filename} (Part {chunk_id})",
            "contents": current_chunk.strip()
        })
    
    return chunks

# Function to generate embeddings
def generate_embedding(text):
    """Generate embeddings for text using OpenAI API"""
    if not OPENAI_API_KEY:
        st.error("OpenAI API key not found. Please make sure it's set in the .env file.")
        return None
    
    openai.api_key = OPENAI_API_KEY
    
    try:
        response = openai.embeddings.create(
            model="text-embedding-ada-002",
            input=text
        )
        return response.data[0].embedding
    except Exception as e:
        st.error(f"Error generating embedding: {str(e)}")
        return None

# Function to test database connection
def test_connection(conn_string):
    try:
        conn = psycopg2.connect(conn_string)
        conn.close()
        return True
    except Exception as e:
        st.error(f"Connection error: {str(e)}")
        return False

# Function to get all tables from the database
def get_tables(conn_string):
    try:
        conn = psycopg2.connect(conn_string)
        cursor = conn.cursor()
        cursor.execute("""
            SELECT table_name FROM information_schema.tables
            WHERE table_schema = 'public'
            ORDER BY table_name;
        """)
        tables = [table[0] for table in cursor.fetchall()]
        cursor.close()
        conn.close()
        return tables
    except Exception as e:
        st.error(f"Error fetching tables: {str(e)}")
        return []

# Function to create a table with vector support
def create_table(conn_string, table_name):
    """Create a new table with vector support"""
    try:
        conn = psycopg2.connect(conn_string)
        cursor = conn.cursor()
        
        # Create pgvector extension if it doesn't exist
        cursor.execute("CREATE EXTENSION IF NOT EXISTS vector;")
        
        # Create table with vector column - use double quotes to preserve case
        cursor.execute(f"""
            CREATE TABLE IF NOT EXISTS "{table_name}" (
                id SERIAL PRIMARY KEY,
                chunk_id VARCHAR(255),
                title VARCHAR(255),
                contents TEXT,
                vector vector(1536)
            )
        """)
        
        conn.commit()
        cursor.close()
        conn.close()
        return True
    except Exception as e:
        st.error(f"Error creating table: {str(e)}")
        return False

# Function to process document and generate embeddings
def process_document(file, file_extension, conn_string, table_name):
    """Process document, split into chunks if needed, and store with embeddings"""
    # First make sure the table exists
    if not ensure_table_exists(conn_string, table_name):
        st.error(f"Error: Table '{table_name}' does not exist and could not be created")
        return False
    
    conn = None
    try:
        # Extract text based on file type
        if file_extension == 'pdf':
            document_text = extract_text_from_pdf(file)
        elif file_extension == 'docx':
            document_text = extract_text_from_docx(file)
        else:
            st.error("Unsupported file type")
            return False
        
        # Count tokens in the entire document
        total_tokens = count_tokens(document_text)
        st.info(f"Document contains approximately {total_tokens} tokens")
        
        # Determine if the document needs to be split
        if total_tokens > MAX_TOKENS:
            st.info("Document exceeds maximum tokens per chunk and will be split")
            chunks = split_document(document_text, file.name.split('.')[0], MAX_TOKENS)
            st.info(f"Document has been split into {len(chunks)} chunks")
        else:
            chunks = [{
                "chunk_id": f"{file.name.split('.')[0]}-1",
                "title": file.name,
                "contents": document_text
            }]
            st.info("Document processed as a single chunk")
        
        # Connect to database
        conn = psycopg2.connect(conn_string)
        cursor = conn.cursor()
        
        # Process each chunk and generate embeddings
        success_count = 0
        for chunk in chunks:
            embedding = generate_embedding(chunk["contents"])
            
            if embedding:
                embedding_list = embedding  # Already a list
                
                try:
                    # Insert data with embedding - use double quotes around table name to preserve case
                    cursor.execute(
                        f'INSERT INTO "{table_name}" (chunk_id, title, contents, vector) VALUES (%s, %s, %s, %s)',
                        (chunk["chunk_id"], chunk["title"], chunk["contents"], embedding_list)
                    )
                    conn.commit()
                    success_count += 1
                except Exception as e:
                    conn.rollback()
                    st.error(f"Error inserting chunk {chunk['chunk_id']}: {str(e)}")
        
        cursor.close()
        
        if success_count > 0:
            st.success(f"Successfully processed and stored {success_count} chunks with embeddings")
            return True
        else:
            st.warning("No chunks were successfully stored")
            return False
    except Exception as e:
        st.error(f"Error processing document: {str(e)}")
        return False
    finally:
        if conn is not None:
            conn.close()

# Process the document
if uploaded_file is not None and connection_string:
    st.info("Document uploaded successfully")
    
    # Extract text based on file type
    file_extension = uploaded_file.name.split('.')[-1].lower()
    if file_extension == 'pdf':
        document_text = extract_text_from_pdf(uploaded_file)
    elif file_extension == 'docx':
        document_text = extract_text_from_docx(uploaded_file)
    
    # Preview extracted text
    with st.expander("Preview Extracted Text"):
        st.text_area("Document Content", document_text, height=200)
    
    # Test connection
    if test_connection(connection_string):
        st.success("Database connection successful")
        
        # Table selection
        table_option = st.radio(
            "Would you like to create a new table or use an existing one?",
            options=["Create new table", "Use existing table"]
        )
        
        if table_option == "Create new table":
            new_table_name = st.text_input("Enter new table name")
            
            if st.button("Create Table and Upload Data") and new_table_name:
                # Create table with vector support
                if create_table(connection_string, new_table_name):
                    # Process document and generate embeddings
                    process_document(
                        uploaded_file, 
                        file_extension, 
                        connection_string, 
                        new_table_name
                    )
        
        else:  # Use existing table
            tables = get_tables(connection_string)
            
            if tables:
                selected_table = st.selectbox("Select an existing table", tables)
                
                if st.button("Upload Data to Selected Table") and selected_table:
                    # Check if the table has columns
                    try:
                        conn = psycopg2.connect(connection_string)
                        cursor = conn.cursor()
                        cursor.execute(f"""
                            SELECT column_name FROM information_schema.columns
                            WHERE table_name = '{selected_table.lower()}'
                            AND table_schema = 'public'
                        """)
                        columns = [col[0] for col in cursor.fetchall()]
                        cursor.close()
                        conn.close()
                        
                        if not columns:
                            st.warning(f"The selected table '{selected_table}' doesn't have any columns. Please select a different table or create a new one.")
                        else:
                            # First check if the table exists and add required columns if needed
                            if ensure_table_has_required_columns(connection_string, selected_table):
                                # Process document with the existing table
                                process_document(
                                    uploaded_file, 
                                    file_extension, 
                                    connection_string, 
                                    selected_table
                                )
                            else:
                                st.error("Could not prepare table structure. Please try creating a new table instead.")
                    except Exception as e:
                        st.error(f"Error checking table columns: {str(e)}")
            else:
                st.warning("No tables found in the database. Please create a new table.") 